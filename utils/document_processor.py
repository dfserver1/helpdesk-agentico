"""
Document processor: loads, parses, and chunks documents for vector indexing.
Supports PDF, DOCX, TXT, MARKDOWN. Uses recursive character splitting with overlap.
"""

import hashlib
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config.settings import get_settings
from config.logging import get_logger
from core.exceptions import DocumentProcessingError

logger = get_logger("document_processor")


@dataclass
class ProcessedDocument:
    """Result of document processing."""
    original_filename: str
    file_type: str
    file_hash: str
    size_bytes: int
    chunks: List[Document]
    metadata: dict = field(default_factory=dict)


class DocumentProcessor:
    """Process uploaded documents into embeddable chunks."""

    def __init__(self, settings=None):
        self.settings = settings or get_settings()
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.settings.CHUNK_SIZE,
            chunk_overlap=self.settings.CHUNK_OVERLAP,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
            length_function=len,
            keep_separator=False,
        )

    def process(self, file_path: str, tenant_id: int = 1, extra_metadata: dict = None) -> ProcessedDocument:
        """
        Process a document file into chunks.

        Args:
            file_path: Path to the uploaded file
            tenant_id: Tenant (organization) scope
            extra_metadata: Additional metadata to attach to chunks

        Returns:
            ProcessedDocument with chunks ready for vector indexing.
        """
        path = Path(file_path)
        if not path.exists():
            raise DocumentProcessingError(f"File not found: {file_path}")

        file_ext = path.suffix.lower().lstrip(".")
        self._validate_extension(file_ext)
        filename = path.name

        logger.info(f"Processing document: {filename}")
        text = self._extract_text(path, file_ext)
        if not text.strip():
            raise DocumentProcessingError(f"Document appears to be empty: {filename}")

        file_hash = hashlib.sha256(
            path.read_bytes()
        ).hexdigest()

        chunks = self._chunk_text(text, path.name, tenant_id, extra_metadata)

        logger.info(f"Processed {filename}: {len(chunks)} chunks")

        return ProcessedDocument(
            original_filename=filename,
            file_type=file_ext,
            file_hash=file_hash,
            size_bytes=path.stat().st_size,
            chunks=chunks,
            metadata={
                "tenant_id": str(tenant_id),
                "filename": path.name,
                **(extra_metadata or {}),
            },
        )

    def _validate_extension(self, ext: str):
        if ext not in self.settings.ALLOWED_EXTENSIONS:
            raise DocumentProcessingError(
                f"Unsupported file type: .{ext}. Allowed: {', '.join(self.settings.ALLOWED_EXTENSIONS)}"
            )

    def _extract_text(self, path: Path, file_ext: str) -> str:
        """Extract raw text from supported document types."""
        if file_ext == "pdf":
            return self._extract_pdf(path)
        elif file_ext == "docx":
            return self._extract_docx(path)
        else:  # txt, md
            return path.read_text(encoding="utf-8", errors="replace")

    def _extract_pdf(self, path: Path) -> str:
        """Extract text from PDF including page numbers."""
        try:
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            parts = []
            for i, page in enumerate(reader.pages, 1):
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(text)
            return "\n\n".join(parts)
        except ImportError:
            raise DocumentProcessingError("PDF support requires pypdf. Install: pip install pypdf")
        except Exception as e:
            raise DocumentProcessingError(f"Failed to read PDF {path.name}: {e}")

    def _extract_docx(self, path: Path) -> str:
        """Extract from DOCX including tables."""
        try:
            import docx

            doc = docx.Document(str(path))
            parts = [para.text for para in doc.paragraphs if para.text.strip()]

            # Extract table content
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        parts.append(" | ".join(filter(None, cells)))

            return "\n".join(parts)
        except ImportError:
            raise DocumentProcessingError("DOCX support requires python-docx. Install: pip install python-docx")
        except Exception as e:
            raise DocumentProcessingError(f"Failed to read DOCX {path.name}: {e}")

    def _chunk_text(
        self,
        text: str,
        filename: str,
        tenant_id: int,
        extra_metadata: dict = None
    ) -> List[Document]:
        """Split text into Document chunks with metadata."""
        raw_chunks = self.text_splitter.split_text(text)

        # document_id lets callers delete all chunks of a file (see
        # vector_db.chroma_store.delete_document_chunks). Falls back to a
        # stable hash of the filename when not supplied explicitly.
        document_id = (extra_metadata or {}).get("document_id") or hashlib.sha256(
            filename.encode("utf-8")
        ).hexdigest()

        base_metadata = {
            "source": filename,
            "tenant_id": str(tenant_id),
            "document_id": document_id,
            **(extra_metadata or {}),
        }

        documents = []
        for i, chunk in enumerate(raw_chunks):
            if not chunk.strip():
                continue
            documents.append(
                Document(
                    page_content=chunk.strip(),
                    metadata={
                        **base_metadata,
                        "chunk_index": i,
                        "chunk_total": len(raw_chunks),
                    },
                )
            )
        return documents


# --- Module-level factory ---
processor = DocumentProcessor()


def process_document(file_path: str, tenant_id: int = 1, extra_metadata: dict = None) -> ProcessedDocument:
    """Convenience factory using the module-level processor."""
    return processor.process(file_path, tenant_id, extra_metadata)